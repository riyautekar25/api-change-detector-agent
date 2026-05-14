import re
from docx import Document

OLD_ORDER_FILE = "old_snapshot/order.py"
NEW_ORDER_FILE = "services/order.py"

LOG_FILE = "logs/api_logs.txt"


def read_file(path):

    with open(path, "r") as file:
        return file.read()


def extract_api_details(code):

    api_details = {}

    # Extract endpoint
    endpoint_match = re.search(
        r'@app\.post\("(.+?)"\)',
        code
    )

    if endpoint_match:
        api_details["endpoint"] = endpoint_match.group(1)

    # Extract function definition
    function_match = re.search(
        r'def\s+(\w+)\((.*?)\):',
        code,
        re.DOTALL
    )

    if function_match:

        api_details["function_name"] = function_match.group(1)

        parameters = function_match.group(2)

        cleaned_params = []

        for param in parameters.split(","):

            param = param.strip()

            if param:
                cleaned_params.append(param)

        api_details["parameters"] = cleaned_params

    return api_details


def compare_apis(old_api, new_api):

    changes = []

    # Compare endpoint
    if old_api["endpoint"] != new_api["endpoint"]:

        changes.append(
            f"Endpoint changed from "
            f"{old_api['endpoint']} "
            f"to "
            f"{new_api['endpoint']}"
        )

    # Compare parameters
    old_params = set(old_api["parameters"])
    new_params = set(new_api["parameters"])

    added_params = new_params - old_params
    removed_params = old_params - new_params

    if added_params:

        changes.append(
            f"New parameters added: "
            f"{', '.join(added_params)}"
        )

    if removed_params:

        changes.append(
            f"Parameters removed: "
            f"{', '.join(removed_params)}"
        )

    return changes


def analyze_logs():

    with open(LOG_FILE, "r") as file:

        logs = file.readlines()

    recent_logs = logs[-5:]

    return recent_logs


def generate_document(changes, logs):

    doc = Document()

    doc.add_heading(
        "OpenAPI Change Detection Report",
        level=1
    )

    doc.add_heading(
        "Detected API Changes",
        level=2
    )

    if changes:

        for change in changes:
            doc.add_paragraph(change)

    else:
        doc.add_paragraph("No API changes detected")

    doc.add_heading(
        "Recent Logs",
        level=2
    )

    for log in logs:
        doc.add_paragraph(log.strip())

    doc.save("openapi_change_report.docx")


def main():

    old_code = read_file(OLD_ORDER_FILE)
    new_code = read_file(NEW_ORDER_FILE)

    old_api = extract_api_details(old_code)
    new_api = extract_api_details(new_code)

    changes = compare_apis(old_api, new_api)

    logs = analyze_logs()

    print("\n========== API CHANGE ANALYSIS ==========\n")

    if changes:

        for change in changes:
            print(change)

    else:
        print("No changes detected")

    print("\nGenerating Word document...\n")

    generate_document(changes, logs)

    print("Document generated successfully")


if __name__ == "__main__":
    main()